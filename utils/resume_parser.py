import pdfplumber
import docx
import re

def extract_text_from_pdf(path):
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(path):
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)

def clean_text(text):
    # Simple cleaning
    text = re.sub(r'\r', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def parse_resume(file_like, filename):
    """
    file_like: a file-like object (Streamlit uploaded file)
    filename: name string (to detect extension)
    Returns cleaned text.
    """
    ext = filename.lower().split('.')[-1]
    if ext == "pdf":
        # pdfplumber needs a path or bytes-like; file_like can provide .read()
        content = file_like.read()
        # write to temporary BytesIO? pdfplumber can open bytes via pdfplumber.open(io.BytesIO(...))
        import io
        return clean_text(extract_text_from_pdf_from_bytes(content))
    elif ext in ("docx", "doc"):
        import io
        return clean_text(extract_text_from_docx_from_bytes(file_like.read()))
    else:
        # fallback - attempt to read raw text
        try:
            txt = file_like.read().decode("utf-8")
            return clean_text(txt)
        except Exception:
            return ""

# helpers for bytes
def extract_text_from_pdf_from_bytes(bytes_data):
    import io
    text = []
    with pdfplumber.open(io.BytesIO(bytes_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx_from_bytes(bytes_data):
    import io, zipfile
    # python-docx can't read from bytes directly, so write to temp file-like.
    from docx import Document
    bio = io.BytesIO(bytes_data)
    doc = Document(bio)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)